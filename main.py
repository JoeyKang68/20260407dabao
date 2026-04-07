# -*- coding: utf-8 -*-
from kivy.app import App
from kivy.uix.button import Button
from kivy.uix.boxlayout import BoxLayout
from kivy.core.text import LabelBase, DEFAULT_FONT
from kivy.uix.image import Image
from kivy.uix.textinput import TextInput
from kivy.uix.spinner import Spinner
import datetime
import docx
import csv
from kivy.lang import Builder

LabelBase.register(DEFAULT_FONT, 'simfang.ttf')
LabelBase.register(name='SIMKAL',fn_regular='simfang.ttf')

# 应用全局样式，解决下拉菜单选项中文乱码问题
Builder.load_string('''
<SpinnerOption>:
    font_name: 'SIMKAL'
''')

class MyApp(App):
    global options1
    global options2
    with open(r"条款.csv", 'r', encoding='gbk') as f:
        rows = list(csv.reader(f))
    
    options1 = [row[0] for row in rows if len(row) > 0 and row[0]]
    options2 = [row[1] for row in rows if len(row) > 1 and row[1]]
    def build(self):

        # 创建一个BoxLayout作为容器
        layout = BoxLayout(orientation='vertical', spacing=1, padding=1)
        #设置背景
        #background_image = Image(source='顺顺荔荔.jpg', allow_stretch=True, keep_ratio=False)
        # 创建多个TextInput实例并添加到布局中
        input1 = TextInput(text='350627'+datetime.datetime.now().strftime('%Y%m%d')+'01',font_name='SIMKAL', multiline=True, size_hint=(0.6,0.13), pos_hint={'x':0.2,'y':0.5})
        input2 = TextInput(text=datetime.datetime.now().strftime('%Y年%m月%d日%H时%M分许'),font_name='SIMKAL', multiline=False, size_hint=(0.6,0.13), pos_hint={'x':0.2,'y':0.5})
        input3 = TextInput(hint_text='天气',font_name='SIMKAL', multiline=True, size_hint=(0.6,0.13), pos_hint={'x':0.2,'y':0.5})
        input4 = TextInput(text='福建省漳州市南靖县', font_name='SIMKAL', multiline=True, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})

        input11 = TextInput(hint_text='甲当事人',font_name='SIMKAL', multiline=True, size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})
        input12 = TextInput(hint_text='甲驾驶证或身份证',font_name='SIMKAL', multiline=True, size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})
        input13 = TextInput(hint_text='甲联系方式',font_name='SIMKAL', multiline=True, size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})
        input14 = TextInput(hint_text='甲交通方式',font_name='SIMKAL', multiline=True, size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})
        input15 = TextInput(text='闽 E',font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})
        input16 = TextInput(text='///////////', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})

        input21 = TextInput(hint_text='乙当事人', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        input22 = TextInput(hint_text='乙驾驶证或身份证', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        input23 = TextInput(hint_text='乙联系方式', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        input24 = TextInput(hint_text='乙交通方式', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        input25 = TextInput(text='闽 E', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        input26 = TextInput(text='///////////', font_name='SIMKAL', multiline=False, size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})

        spinner1 = Spinner(values=options1, text='条款1',font_name='SIMKAL',size_hint=(0.6, 0.13),pos_hint={'x': 0.2, 'y': 0.5})
        spinner2 = Spinner(values=options1, text='条款2', font_name='SIMKAL',size_hint=(0.6, 0.13), pos_hint={'x': 0.2, 'y': 0.5})

        input31 = TextInput(hint_text='事故事实', font_name='SIMKAL', multiline=True, size_hint=(0.8, 0.52),pos_hint={'x': 0.1, 'y': 0.5})
        input32 = TextInput(hint_text='事故成因以及事故责任', font_name='SIMKAL', multiline=True,size_hint=(0.8, 0.52), pos_hint={'x': 0.1, 'y': 0.5})
        input33 = TextInput(hint_text='违规条款', font_name='SIMKAL', multiline=True, size_hint=(0.8, 0.52),pos_hint={'x': 0.1, 'y': 0.5})
        list1=[1,'Time','weather','location','Aparty','AID','AContact','Atransportation','Agrade','AInsurance','Bparty','BID','BContact','Btransportation','Bgrade','BInsurance','1ad','2ad','aaaa','bbbb','cccc']
        list2=[input1,input2,input3,input4,input11,input12,input13,input14,input15,input16,input21,input22,input23,input24,input25,input26,spinner1,spinner2,input31,input32,input33]
        #生成事故文本
        def button_callback1(instance):
            a=''
            if spinner1.text !='条款1' and spinner2.text !='条款2':
                a=f'{spintext1}、{spintext2}'
            elif spinner1.text !='条款1' and spinner2.text =='条款2':
                a = f'{spintext1}'
            elif spinner1.text =='条款1' and spinner2.text !='条款2':
                a = f'{spintext2}'
            input31.text = f'''    上述时间、地点:{input11.text}驾驶{input15.text}号{input14.text}沿荆江路由中国银行往实验中学方向行驶至事故地点右转弯时，与十字路口对向左转弯路未按灯通行{input21.text}驾{input25.text}号{input24.text}发生碰撞，造成{input21.text}受伤、车辆受损的交通事故。'''
            input32.text = f'''    本事故中{input21.text}驾车上道路行驶时，未按交通指示灯行驶，其行为是导致本事故发生的主要过错;本事故中{input11.text}驾车上道路行驶，未及时注意路面交通情况、遇况采取措施不及，致事故发生，其行为是导致本事故发生的次要原因。'''
            input33.text = f'''    依照{a}作出责任认定:{input21.text}承担本事故主要责任;{input11.text}承担本事故次要责任。'''
        btn1 = Button(text='生成', font_name='SIMKAL', size_hint=(0.2, 0.1), pos_hint={'x': 0.4, 'y': 0.5},on_press=button_callback1)

        # 定义一个函数，当按钮被点击时调用
        def button_callback2(instance):
            def doc_table_replace(doc_file):
                document=docx.Document(doc_file)
                for para in document.paragraphs:
                    for run in para.runs:
                        if 'num' == run.text:
                            run.text=run.text.replace('num',input1.text)
                #对表格内容进行更改
                for i in range(1, len(list1)):
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text in 'aaaabbbbcccc':
                                    cell.text=cell.text.replace(list1[i],'\t\t\t\t'+list2[i].text)
                                else:
                                    cell.text = cell.text.replace(list1[i], list2[i].text)

                document.save(f"道路交通事故认定书{datetime.datetime.now().strftime('%Y年%m月%d日%H时%M分许')}.docx")
            doc_file= '简易空白test.docx'
            doc_table_replace(doc_file)
            print('替换完成')
        btn2 = Button(text='保存',font_name='SIMKAL', size_hint=(0.2, 0.1),pos_hint={'x': 0.4, 'y': 0.5}, on_press=button_callback2)
        # 绑定选择事件
        spinner1.bind(text=self.on_spinner_select1)
        spinner2.bind(text=self.on_spinner_select2)
        # 将文本输入框添加到布局中
        # 添加背景图片

        #layout.add_widget(background_image)
        for i in list2:
            layout.add_widget(i)

        layout.add_widget(btn1)
        layout.add_widget(btn2)
        return layout

    def on_spinner_select1(self, instance, value):
        global spintext1
        spintext1=options2[options1.index(value)]
    def on_spinner_select2(self, instance, value):
        global spintext2
        spintext2=options2[options1.index(value)]

if __name__ == '__main__':
    MyApp().run()